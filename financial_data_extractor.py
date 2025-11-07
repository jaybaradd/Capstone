# financial_data_extractor.py
import os, time, re, fitz, json, hashlib
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from dotenv import load_dotenv
import pandas as pd

# third-party libs
from groq import Groq                  # Groq API client
import tiktoken                        # token counter for estimation

# env / keys
load_dotenv(override=True)
GROQ_API_KEY = os.getenv("ChatGroq_API_KEY")
if not GROQ_API_KEY:
    raise EnvironmentError("GROQ_API_KEY missing in .env!")

# Initialize Groq client
groq_client = Groq(api_key=GROQ_API_KEY)
TOK_COUNTER = tiktoken.encoding_for_model("gpt-4o-mini")

# financial metrics to extract
FINANCIAL_METRICS = [
    # Revenue & Income
    "Total Revenue", "Net Revenue", "Revenue from Operations", "Total Income",
    "Operating Revenue", "Gross Revenue", "Sales", "Turnover",
    
    # Profit Metrics
    "Net Profit", "Net Income", "Profit After Tax", "PAT",
    "Gross Profit", "Operating Profit", "EBITDA", "EBIT",
    "Profit Before Tax", "PBT", "Operating Income",
    
    # Margins
    "Gross Margin", "Operating Margin", "Net Margin", "EBITDA Margin",
    "Profit Margin", "Operating Profit Margin",
    
    # Cash Flow
    "Operating Cash Flow", "Free Cash Flow", "Cash Flow from Operations",
    "Net Cash Flow", "Cash and Cash Equivalents",
    
    # Balance Sheet
    "Total Assets", "Total Liabilities", "Net Worth", "Shareholders Equity",
    "Total Debt", "Net Debt", "Working Capital", "Current Assets",
    "Current Liabilities", "Non-Current Assets", "Non-Current Liabilities",
    
    # Ratios
    "Current Ratio", "Debt to Equity", "Return on Assets", "ROA",
    "Return on Equity", "ROE", "Return on Capital Employed", "ROCE",
    "Debt Service Coverage Ratio", "Interest Coverage Ratio",
    
    # Per Share Data
    "Earnings Per Share", "EPS", "Book Value Per Share", "BVPS",
    "Dividend Per Share", "DPS", "Net Asset Value Per Share",
    
    # Other Key Metrics
    "Market Capitalization", "Enterprise Value", "Revenue Growth",
    "Profit Growth", "Asset Turnover", "Inventory Turnover"
]

# helper functions
def ntoks(txt: str) -> int:
    """Rough token count estimation"""
    return len(TOK_COUNTER.encode(txt))

def safe_groq_call(prompt: str, max_tokens: int = 1000, temperature: float = 0.2, 
                   tries: int = 3, pause: int = 2) -> str:
    """Safe Groq API call with retry logic"""
    for attempt in range(tries):
        try:
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
                temperature=temperature,
                stream=False
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Groq API error (attempt {attempt+1}/{tries}): {e}")
            if attempt < tries - 1:
                time.sleep(pause * (2 ** attempt))
            else:
                print("All Groq API attempts failed")
                return ""
    return ""

def extract_numbers_from_text(text: str) -> List[Tuple[str, str]]:
    """Extract financial numbers from text using regex patterns"""
    patterns = [
        r'([\d,]+\.?\d*)\s*crores?',
        r'([\d,]+\.?\d*)\s*lakhs?',
        r'([\d,]+\.?\d*)\s*millions?',
        r'([\d,]+\.?\d*)\s*billions?',
        r'₹\s*([\d,]+\.?\d*)',
        r'Rs\.?\s*([\d,]+\.?\d*)',
        r'INR\s*([\d,]+\.?\d*)',
        r'([\d,]+\.?\d*)\s*%',
        r'([\d,]+\.?\d*)\s*percent'
    ]
    
    numbers = []
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            number = match.group(1) if match.group(1) else match.group(0)
            context = text[max(0, match.start()-50):match.end()+50]
            numbers.append((number, context))
    
    return numbers

def standardize_number(num_str: str, unit: str = "") -> float:
    """Convert number string to standardized float value"""
    try:
        # Remove commas and convert to float
        clean_num = float(num_str.replace(',', ''))
        
        # Convert to crores for standardization
        if 'lakh' in unit.lower():
            return clean_num / 100  # Convert lakhs to crores
        elif 'million' in unit.lower():
            return clean_num / 10   # Convert millions to crores
        elif 'billion' in unit.lower():
            return clean_num * 100  # Convert billions to crores
        else:
            return clean_num  # Assume already in crores or appropriate unit
    except:
        return 0.0

class FinancialDataExtractor:
    def __init__(self, company_name: str, default_year: str = None):
        self.company_name = company_name
        self.slug = re.sub(r"[^\w]+", "_", company_name.lower())
        self.reports_data = {}  # Store extracted data by year/quarter
        self.audit_data = {}
        self.default_year = default_year or time.strftime('%Y')  # Default to current year
        
        # Runtime stats
        self.tokens_used = 0
        self.processing_stats = {
            "reports_processed": 0,
            "api_calls": 0,
            "metrics_extracted": 0,
            "years_analyzed": 0,
            "chunks_processed": 0
        }
        
        # Create cache directory
        self.cache_dir = Path(f".cache_{self.slug}")
        self.cache_dir.mkdir(exist_ok=True)
        
        # Optional: Manual filename to period mapping
        self.filename_mapping = {}

    def set_filename_mapping(self, mapping: Dict[str, str]):
        """Set manual filename to period mapping"""
        self.filename_mapping = mapping
        print(f"Set manual filename mapping for {len(mapping)} files")

    def set_report_files(self, files: List[str]):
        """Set report files to process"""
        self.report_files = [f for f in files if os.path.exists(f)]
        print(f"Found {len(self.report_files)} report files")

    def set_audit_files(self, files: List[str]):
        """Set audit files to process"""
        self.audit_files = [f for f in files if os.path.exists(f)]
        print(f"Found {len(self.audit_files)} audit files")

    def _normalize_metric_name(self, metric_name: str) -> str:
        """Normalize metric names for consistent comparison across years"""
        # Convert to lowercase
        normalized = metric_name.lower()
        
        # Remove common variations in parentheses and units
        normalized = re.sub(r'\(.*?\)', '', normalized)  # Remove parentheses content
        normalized = re.sub(r'[₹$]', '', normalized)      # Remove currency symbols
        normalized = re.sub(r'\s*rs\.?\s*', '', normalized)  # Remove Rs.
        normalized = re.sub(r'\s*inr\s*', '', normalized)    # Remove INR
        normalized = re.sub(r'\s*crores?\s*', '', normalized)  # Remove crores
        normalized = re.sub(r'\s*lakhs?\s*', '', normalized)   # Remove lakhs
        normalized = re.sub(r'\s*millions?\s*', '', normalized) # Remove millions
        normalized = re.sub(r'\s*billions?\s*', '', normalized) # Remove billions
        normalized = re.sub(r'\s+', ' ', normalized).strip()  # Normalize whitespace
        
        # Common abbreviations standardization
        normalized = normalized.replace('profit after tax', 'pat')
        normalized = normalized.replace('profit before tax', 'pbt')
        normalized = normalized.replace('earnings per share', 'eps')
        normalized = normalized.replace('return on equity', 'roe')
        normalized = normalized.replace('return on assets', 'roa')
        normalized = normalized.replace('return on capital employed', 'roce')
        normalized = normalized.replace('book value per share', 'bvps')
        normalized = normalized.replace('dividend per share', 'dps')
        normalized = normalized.replace('earnings before interest tax depreciation amortization', 'ebitda')
        normalized = normalized.replace('earnings before interest tax', 'ebit')
        
        return normalized

    def run(self):
        """Main processing pipeline"""
        start_time = time.perf_counter()
        print(f"Starting financial data extraction for {self.company_name}")

        try:
            # Step 1: Extract financial data from all reports
            self._extract_financial_data()
            
            # Step 2: Extract audit information
            audit_summary = self._extract_audit_info()
            
            # Step 3: Create comparison tables
            comparison_tables = self._create_comparison_tables()
            
            # Step 4: Generate insights and analysis
            insights = self._generate_insights()
            
            # Step 5: Create comprehensive report
            self._create_final_report(comparison_tables, audit_summary, insights)
            
            self._print_stats(time.perf_counter() - start_time)
            
        except Exception as e:
            print(f"Pipeline error: {e}")
            raise

    def _extract_financial_data(self):
        """Extract financial data from all report files"""
        print("\nExtracting financial data from reports...")
        
        for i, file_path in enumerate(self.report_files):
            print(f"\n  Processing file {i+1}/{len(self.report_files)}: {os.path.basename(file_path)}")
            
            # Extract text from PDF
            text = self._extract_text_from_pdf(file_path)
            if not text:
                print(f"  No text extracted from {file_path}")
                continue
            
            # Identify time period (year/quarter)
            period = self._identify_period(text, file_path)
            print(f"    Identified period: {period}")
            
            # Extract financial metrics
            metrics = self._extract_metrics_from_text(text, period)
            
            if metrics:
                self.reports_data[period] = metrics
                self.processing_stats["reports_processed"] += 1
                self.processing_stats["metrics_extracted"] += len(metrics)
                print(f"    Extracted {len(metrics)} financial metrics")
            else:
                print(f"    No financial metrics extracted")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            return text
        except Exception as e:
            print(f"  Error extracting text from {file_path}: {e}")
            return ""

    def _identify_period(self, text: str, file_path: str) -> str:
        """Identify the time period from text and filename"""
        filename = os.path.basename(file_path).lower()
        
        # Check manual mapping first
        if filename in self.filename_mapping:
            return self.filename_mapping[filename]
        
        # Look for year patterns
        year_patterns = [
            r'(20\d{2})',
            r'fy(\d{2})',
            r'(\d{4})'
        ]
        
        # Look for quarter patterns
        quarter_patterns = [
            r'q([1-4])',
            r'quarter\s*([1-4])',
            r'(first|second|third|fourth)\s*quarter'
        ]
        
        found_year = None
        found_quarter = None
        
        # Check filename first
        for pattern in year_patterns:
            match = re.search(pattern, filename)
            if match:
                year_str = match.group(1)
                if len(year_str) == 2:
                    found_year = f"20{year_str}"
                else:
                    found_year = year_str
                break
        
        for pattern in quarter_patterns:
            match = re.search(pattern, filename)
            if match:
                quarter_str = match.group(1)
                if quarter_str.isdigit():
                    found_quarter = f"Q{quarter_str}"
                else:
                    quarter_map = {"first": "Q1", "second": "Q2", "third": "Q3", "fourth": "Q4"}
                    found_quarter = quarter_map.get(quarter_str.lower())
                break
        
        # More intelligent text search if no year found in filename
        if not found_year:
            # Look for specific reporting period patterns first
            text_search = text[:5000].lower()  # Search first 5000 chars
            
            # Look for patterns that indicate the actual reporting period
            period_patterns = [
                r'(?:for the year ended|year ended).*?(20\d{2})',
                r'(?:for the quarter ended|quarter ended).*?(20\d{2})',
                r'(?:for the three months ended|three months ended).*?(20\d{2})',
                r'(?:for the six months ended|six months ended).*?(20\d{2})',
                r'(?:for the nine months ended|nine months ended).*?(20\d{2})',
                r'(?:reporting period|period ended).*?(20\d{2})',
                r'(?:financial year|fy).*?(20\d{2})',
                r'(?:ended|ending).*?(20\d{2})'
            ]
            
            for pattern in period_patterns:
                match = re.search(pattern, text_search)
                if match:
                    found_year = match.group(1)
                    print(f"    Found year from period pattern: {found_year}")
                    break
            
            # If still not found, use default year instead of random text search
            if not found_year:
                found_year = self.default_year
                print(f"    No year found in filename '{filename}', defaulting to {found_year}")
        
        # Look for quarter info in text if not found in filename
        if not found_quarter:
            text_lower = text.lower()
            quarter_match = re.search(r'(?:quarter|q)\s*([1-4])', text[:3000])
            if quarter_match:
                q_num = quarter_match.group(1)
                found_quarter = f"Q{q_num}"
        
        # Format period string
        if found_year and found_quarter:
            return f"{found_year} {found_quarter}"
        elif found_year:
            return f"{found_year} Annual"
        else:
            return f"Unknown Period - {filename}"

    def _extract_metrics_from_text(self, text: str, period: str) -> Dict[str, float]:
        """Extract financial metrics from text using AI with comprehensive document coverage"""
        print(f"    Extracting metrics for {period}...")
        
        # Original chunking parameters
        chunk_size = 25000  # Larger chunks
        overlap = 3000      # More overlap for better context
        
        # Create overlapping chunks
        chunks = []
        text_len = len(text)
        i = 0
        
        while i < text_len:
            end_pos = min(i + chunk_size, text_len)
            chunk = text[i:end_pos]
            chunks.append(chunk)
            i += chunk_size - overlap
            
            # Stop if we've covered the entire text
            if end_pos >= text_len:
                break
        
        print(f"      Processing {len(chunks)} chunks")
        
        all_metrics = {}
        metric_original_names = {}  # Track original names for display
        
        # Process chunks with original delays
        for i, chunk in enumerate(chunks):
            print(f"      Processing chunk {i+1}/{len(chunks)}")
            
            # Add delay between API calls
            if i > 0:
                time.sleep(2)
            
            prompt = f"""
Extract specific financial metrics from this financial document text for {period}.

Look for these key metrics and their values:
- Revenue/Sales figures (Total Revenue, Net Revenue, Operating Revenue, Sales, Turnover)
- Profit figures (Net Profit, Net Income, Gross Profit, Operating Profit, EBITDA, EBIT, PAT, PBT)
- Margin percentages (Gross Margin, Net Margin, EBITDA Margin, Operating Margin, Profit Margin)
- Cash flow figures (Operating Cash Flow, Free Cash Flow, Cash Flow from Operations, Net Cash Flow)
- Balance sheet items (Total Assets, Total Liabilities, Total Debt, Net Debt, Net Worth, Shareholders Equity, Working Capital, Current Assets, Current Liabilities)
- Ratios (Current Ratio, Debt to Equity, ROE, ROA, ROCE, Interest Coverage Ratio, Debt Service Coverage Ratio)
- Per share data (EPS, Earnings Per Share, Book Value Per Share, BVPS, Dividend Per Share, DPS)
- Growth metrics (Revenue Growth, Profit Growth)
- Other metrics (Market Capitalization, Enterprise Value, Asset Turnover, Inventory Turnover)

Document text:
{chunk}

Return ONLY a JSON object with metric names as keys and numerical values as values. 
For percentages, include the % symbol in the key name.
For currency amounts, convert to crores and include the unit in the key name.
Use null for metrics not found.
Be thorough and extract ALL available metrics from this text chunk.

Example format:
{{
    "Total Revenue (₹ Crores)": 15234.5,
    "Net Profit (₹ Crores)": 2845.2,
    "EBITDA Margin (%)": 18.5,
    "EPS (₹)": 45.2,
    "Debt to Equity Ratio": 0.65,
    "Operating Cash Flow (₹ Crores)": 3245.8,
    "Current Ratio": 1.25,
    "ROE (%)": 15.2
}}

JSON:
"""
            
            response = safe_groq_call(prompt, max_tokens=1500, temperature=0.1)
            self.processing_stats["api_calls"] += 1
            self.processing_stats["chunks_processed"] += 1
            self.tokens_used += ntoks(prompt) + 1500
            
            if response:
                try:
                    # Clean response to extract JSON
                    json_match = re.search(r'\{.*\}', response, re.DOTALL)
                    if json_match:
                        json_str = json_match.group(0)
                        chunk_metrics = json.loads(json_str)
                        
                        # Store with normalized keys while preserving original display names
                        for metric, value in chunk_metrics.items():
                            if value is not None and value != 0:
                                normalized_key = self._normalize_metric_name(metric)
                                
                                # Keep original name for display (use first occurrence or more descriptive one)
                                if normalized_key not in metric_original_names:
                                    metric_original_names[normalized_key] = metric
                                elif len(metric) > len(metric_original_names[normalized_key]):
                                    # Prefer longer/more descriptive names
                                    metric_original_names[normalized_key] = metric
                                
                                # Store value with normalized key
                                if normalized_key not in all_metrics:
                                    all_metrics[normalized_key] = value
                                else:
                                    # If duplicate, prefer non-zero values
                                    if value != 0:
                                        all_metrics[normalized_key] = value
                                
                except json.JSONDecodeError as e:
                    print(f"      JSON decode error: {e}")
                    continue
                except Exception as e:
                    print(f"      Error processing chunk response: {e}")
                    continue
        
        print(f"      Total unique metrics extracted: {len(all_metrics)}")
        
        # Store in reports_data with display names but use normalized keys internally
        # We'll return a dict with display names as keys for backward compatibility
        display_metrics = {}
        for normalized_key, value in all_metrics.items():
            display_name = metric_original_names.get(normalized_key, normalized_key.title())
            display_metrics[display_name] = value
        
        # Also store the normalized mapping for this period
        if not hasattr(self, 'normalized_metrics'):
            self.normalized_metrics = {}
        if period not in self.normalized_metrics:
            self.normalized_metrics[period] = {}
        
        for normalized_key, value in all_metrics.items():
            display_name = metric_original_names.get(normalized_key, normalized_key.title())
            self.normalized_metrics[period][normalized_key] = {
                'value': value,
                'display_name': display_name
            }
        
        return display_metrics

    def _extract_audit_info(self) -> str:
        """Extract audit information from audit files"""
        if not hasattr(self, 'audit_files') or not self.audit_files:
            return "No audit files provided for analysis."
        
        print("\nExtracting audit information...")
        
        combined_text = ""
        for file_path in self.audit_files:
            text = self._extract_text_from_pdf(file_path)
            if text:
                combined_text += f"\n--- {os.path.basename(file_path)} ---\n{text}\n"
        
        if not combined_text.strip():
            return "Could not extract text from audit files."
        
        # Process audit information
        prompt = f"""
Extract comprehensive auditor information from these financial documents:

{combined_text[:40000]}  # Limit to first 40k characters

Extract:
1. Current and previous auditor names
2. Audit fees for each year (if available)
3. Any changes in auditors and reasons
4. Audit opinions and any qualifications
5. Non-audit services provided
6. Key audit matters or concerns

Provide a structured summary in paragraph form suitable for investors.
Maximum 200 words.
"""
        
        response = safe_groq_call(prompt, max_tokens=300, temperature=0.2)
        self.processing_stats["api_calls"] += 1
        self.tokens_used += ntoks(prompt) + 300
        
        return response or "Unable to extract audit information."

    def _create_comparison_tables(self) -> Dict[str, pd.DataFrame]:
        """Create comparison tables from extracted data using normalized metric names"""
        print("\nCreating comparison tables...")
        
        if not self.reports_data:
            print("  No financial data available for comparison")
            return {}
        
        # Debug: Print what we have
        print(f"  Total periods with data: {len(self.reports_data)}")
        for period in self.reports_data:
            print(f"    {period}: {len(self.reports_data[period])} metrics")
        
        # Use normalized metrics for matching across periods
        if not hasattr(self, 'normalized_metrics') or not self.normalized_metrics:
            print("  Warning: No normalized metrics available, falling back to direct matching")
            # Fallback to old behavior
            metrics_by_period = {}
            all_metrics = set()
            display_names = {}
            
            for period, metrics in self.reports_data.items():
                for metric, value in metrics.items():
                    all_metrics.add(metric)
                    display_names[metric] = metric
                    if metric not in metrics_by_period:
                        metrics_by_period[metric] = {}
                    metrics_by_period[metric][period] = value
        else:
            # Organize data by NORMALIZED metric name
            metrics_by_period = {}
            all_metrics = set()
            display_names = {}  # Map normalized -> best display name
            
            for period, normalized_data in self.normalized_metrics.items():
                for normalized_key, data in normalized_data.items():
                    all_metrics.add(normalized_key)
                    value = data['value']
                    display_name = data['display_name']
                    
                    # Track best display name (prefer more descriptive/longer ones)
                    if normalized_key not in display_names:
                        display_names[normalized_key] = display_name
                    elif len(display_name) > len(display_names[normalized_key]):
                        display_names[normalized_key] = display_name
                    
                    if normalized_key not in metrics_by_period:
                        metrics_by_period[normalized_key] = {}
                    metrics_by_period[normalized_key][period] = value
        
        print(f"  Total unique metrics (after normalization): {len(all_metrics)}")
        print(f"  Sample normalized metrics: {list(all_metrics)[:10]}")
        
        # Create comparison tables by category
        tables = {}
        
        # Revenue and Profitability
        revenue_metrics = [m for m in all_metrics if any(term in m.lower() for term in 
                          ['revenue', 'sales', 'turnover', 'income', 'profit', 'ebitda', 'ebit', 'pat', 'pbt'])]
        print(f"  Revenue metrics candidates: {len(revenue_metrics)}")
        if revenue_metrics:
            df = self._create_table(revenue_metrics, metrics_by_period, display_names, 'Revenue & Profitability')
            if not df.empty:
                tables['Revenue & Profitability'] = df
                print(f"    Created Revenue & Profitability table with {len(df)} rows")
        
        # Margins and Ratios
        margin_metrics = [m for m in all_metrics if any(term in m.lower() for term in 
                         ['margin', 'ratio', 'roe', 'roa', 'roce', '%', 'percent'])]
        print(f"  Margin metrics candidates: {len(margin_metrics)}")
        if margin_metrics:
            df = self._create_table(margin_metrics, metrics_by_period, display_names, 'Margins & Ratios')
            if not df.empty:
                tables['Margins & Ratios'] = df
                print(f"    Created Margins & Ratios table with {len(df)} rows")
        
        # Balance Sheet
        balance_metrics = [m for m in all_metrics if any(term in m.lower() for term in 
                          ['assets', 'liabilities', 'equity', 'debt', 'cash', 'working capital'])]
        print(f"  Balance sheet metrics candidates: {len(balance_metrics)}")
        if balance_metrics:
            df = self._create_table(balance_metrics, metrics_by_period, display_names, 'Balance Sheet')
            if not df.empty:
                tables['Balance Sheet'] = df
                print(f"    Created Balance Sheet table with {len(df)} rows")
        
        # Per Share Data
        per_share_metrics = [m for m in all_metrics if any(term in m.lower() for term in 
                            ['eps', 'per share', 'dividend', 'book value', 'bvps', 'dps'])]
        print(f"  Per share metrics candidates: {len(per_share_metrics)}")
        if per_share_metrics:
            df = self._create_table(per_share_metrics, metrics_by_period, display_names, 'Per Share Data')
            if not df.empty:
                tables['Per Share Data'] = df
                print(f"    Created Per Share Data table with {len(df)} rows")
        
        # Cash Flow
        cash_metrics = [m for m in all_metrics if any(term in m.lower() for term in 
                       ['cash flow', 'operating cash', 'free cash'])]
        print(f"  Cash flow metrics candidates: {len(cash_metrics)}")
        if cash_metrics:
            df = self._create_table(cash_metrics, metrics_by_period, display_names, 'Cash Flow')
            if not df.empty:
                tables['Cash Flow'] = df
                print(f"    Created Cash Flow table with {len(df)} rows")
        
        print(f"  Final tables created: {len(tables)}")
        return tables

    def _create_table(self, metrics: List[str], metrics_by_period: Dict, 
                      display_names: Dict = None, table_name: str = "") -> pd.DataFrame:
        """Create a pandas DataFrame for specific metrics using normalized names with more lenient filtering"""
        print(f"    Creating table: {table_name}")
        
        if display_names is None:
            display_names = {m: m for m in metrics}
        
        # Get all periods
        all_periods = set()
        for metric in metrics:
            if metric in metrics_by_period:
                all_periods.update(metrics_by_period[metric].keys())
        
        # Filter to only annual periods for completeness check
        annual_periods = [p for p in all_periods if 'annual' in p.lower()]
        print(f"      Annual periods found: {len(annual_periods)} - {annual_periods}")
        
        # Sort periods chronologically
        def sort_periods(periods):
            """Sort periods chronologically"""
            def period_key(period):
                # Extract year and quarter for sorting
                year_match = re.search(r'(\d{4})', period)
                quarter_match = re.search(r'Q(\d)', period)
                
                year = int(year_match.group(1)) if year_match else 9999
                quarter = int(quarter_match.group(1)) if quarter_match else 0
                
                # Annual reports get quarter 0 for sorting
                if 'annual' in period.lower():
                    quarter = 0
                
                return (year, quarter)
            
            return sorted(periods, key=period_key)
        
        sorted_periods = sort_periods(list(all_periods))
        print(f"      All periods (sorted): {sorted_periods}")
        
        # More lenient filtering: require at least 2 annual periods (not all 3)
        # and allow some missing values
        filtered_metrics = []
        for metric in metrics:
            if metric in metrics_by_period:
                metric_periods = set(metrics_by_period[metric].keys())
                
                # Count how many annual periods have data
                annual_data_count = sum(1 for p in annual_periods if p in metric_periods)
                
                # Require at least 2 annual periods with data (instead of all)
                if annual_data_count >= min(2, len(annual_periods)):
                    # Check if at least some values are valid
                    valid_value_count = 0
                    for period in annual_periods:
                        value = metrics_by_period[metric].get(period)
                        # More lenient: accept any numeric value including 0
                        if value is not None and value != "N/A":
                            valid_value_count += 1
                    
                    # Require at least 2 valid values
                    if valid_value_count >= min(2, len(annual_periods)):
                        filtered_metrics.append(metric)
        
        print(f"      Metrics after filtering: {len(filtered_metrics)}/{len(metrics)}")
        
        # If no metrics pass the filter, try even more lenient approach
        if not filtered_metrics and len(annual_periods) > 0:
            print(f"      No metrics passed strict filter, using lenient filter...")
            for metric in metrics:
                if metric in metrics_by_period:
                    # Just require ANY data point with a valid value
                    if any(v is not None and v != "N/A" 
                           for v in metrics_by_period[metric].values()):
                        filtered_metrics.append(metric)
            print(f"      Metrics after lenient filtering: {len(filtered_metrics)}/{len(metrics)}")
        
        # Create DataFrame with display names
        data = []
        for metric in filtered_metrics:
            if metric in metrics_by_period:
                # Use display name for the metric column
                display_name = display_names.get(metric, metric.title())
                row = [display_name]
                for period in sorted_periods:
                    value = metrics_by_period[metric].get(period, "N/A")
                    row.append(value)
                data.append(row)
        
        if not data:
            print(f"      WARNING: No data rows created for {table_name}")
            return pd.DataFrame()
        
        columns = ['Metric'] + sorted_periods
        df = pd.DataFrame(data, columns=columns)
        print(f"      Created table with {len(df)} rows and {len(df.columns)} columns")
        return df

    def _generate_insights(self) -> str:
        """Generate insights from the comparison data"""
        print("\nGenerating insights...")
        
        if not self.reports_data:
            return "No data available for insights generation."
        
        # Prepare data summary for AI analysis
        data_summary = ""
        for period, metrics in self.reports_data.items():
            data_summary += f"\n{period}:\n"
            for metric, value in metrics.items():
                data_summary += f"  {metric}: {value}\n"
        
        prompt = f"""
Analyze this multi-year financial data and generate key insights for investors:

{data_summary}

Provide analysis on:
1. Revenue growth trends
2. Profitability improvements or deteriorations
3. Margin analysis
4. Balance sheet strength
5. Cash flow patterns
6. Key performance indicators trends
7. Areas of concern or strength
8. Year-over-year comparisons

Structure the insights in a clear, investor-friendly format.
Maximum 500 words.
"""
        
        response = safe_groq_call(prompt, max_tokens=800, temperature=0.3)
        self.processing_stats["api_calls"] += 1
        self.tokens_used += ntoks(prompt) + 800
        
        return response or "Unable to generate insights."

    def _create_final_report(self, tables: Dict[str, pd.DataFrame], audit_summary: str, insights: str):
        """Create the final Word document report"""
        print("\nCreating final report...")
        
        doc = Document()
        
        # Title
        doc.add_heading(f"{self.company_name} - Financial Analysis Report", 0)
        doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Periods Analyzed: {len(self.reports_data)}")
        doc.add_paragraph(f"Default Year Used: {self.default_year}")
        
        # Executive Summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(insights)
        
        # Financial Comparison Tables
        doc.add_heading("Financial Comparison Tables", 1)
        
        for table_name, df in tables.items():
            if df.empty:
                continue
                
            doc.add_heading(table_name, 2)
            
            # Add table to document
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                header_cells[i].text = str(column)
            
            # Data rows
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            
            doc.add_paragraph("")  # Add space after table
        
        # Audit Information
        doc.add_heading("Audit Information", 1)
        doc.add_paragraph(audit_summary)
        
        # Processing Statistics
        doc.add_heading("Processing Statistics", 1)
        stats_text = f"""
Reports Processed: {self.processing_stats['reports_processed']}
Total Metrics Extracted: {self.processing_stats['metrics_extracted']}
API Calls Made: {self.processing_stats['api_calls']}
Chunks Processed: {self.processing_stats['chunks_processed']}
Estimated Tokens Used: {self.tokens_used:,}
Data Points: {len(self.reports_data)}
Default Year Used: {self.default_year}
"""
        doc.add_paragraph(stats_text)
        
        # Save document
        output_path = f"{self.slug}_financial_analysis.docx"
        try:
            doc.save(output_path)
            print(f"Report saved: {output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")

    def _print_stats(self, runtime: float):
        """Print processing statistics"""
        print(f"\n{'='*60}")
        print(f"Processing Statistics for {self.company_name}")
        print(f"{'='*60}")
        print(f"Total Runtime: {runtime:.1f}s")
        print(f"Reports Processed: {self.processing_stats['reports_processed']}")
        print(f"Metrics Extracted: {self.processing_stats['metrics_extracted']}")
        print(f"API Calls: {self.processing_stats['api_calls']}")
        print(f"Chunks Processed: {self.processing_stats['chunks_processed']}")
        print(f"Tokens Used: {self.tokens_used:,}")
        print(f"Data Points: {len(self.reports_data)}")
        print(f"Default Year: {self.default_year}")
        print(f"{'='*60}")

# Demo Usage
if __name__ == "__main__":
    # Example usage with improved period identification
    extractor = FinancialDataExtractor("Infosys2025", default_year="2025")
    
    # Optional: Set manual filename mapping for better control
    extractor.set_filename_mapping({
        "q2-2025.pdf": "2025 Q2",
        "q3-2025.pdf": "2025 Q3",
        "infosys-ar-25.pdf": "2025 Annual",
        "infosys-ar-24.pdf": "2024 Annual", 
        "infosys-ar-23.pdf": "2023 Annual"
    })
    
    # Set report files (multiple years/quarters)
    extractor.set_report_files([
        "infosys-ar-25.pdf",  # Annual Report 2025
        "infosys-ar-24.pdf",  # Annual Report 2024
        "infosys-ar-23.pdf",  # Annual Report 2023
        "q3-2025.pdf",    # Q3 2025
        "q2-2025.pdf",    # Q2 2025
    ])
    
    # Set audit files
    extractor.set_audit_files([
        "infosys-ar-25.pdf",
        "infosys-ar-24.pdf",
        "infosys-ar-23.pdf"
    ])
    
    # Run analysis
    extractor.run()
    
    print("\nFinancial analysis complete!")
    print("Check the generated Word document for comparison tables and insights.")
    print("The tool has extracted financial metrics and created year-over-year comparisons.")
    print("Period identification has been improved to avoid incorrect year detection.")